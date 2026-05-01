"""Generate a Word document containing all references for the H-1B paper.

Outputs: outputs/paper/References_Xiao Xu_Cassie Zhang.docx

Citations are organized by category:
  1. Primary data sources
  2. Methodology / software
  3. H-1B policy and labor-market literature
  4. Government & think-tank reports
  5. Classification systems & technical standards
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUT_PATH = Path(
    r"D:/Georgetown/26spring_semester/Massive_Data/project2/project2/"
    r"outputs/paper/References_Xiao Xu_Cassie Zhang.docx"
)


REFERENCES = {
    "1. Primary Data Sources": [
        (
            "U.S. Citizenship and Immigration Services (USCIS). (2024). "
            "H-1B Employer Data Hub — Fiscal Year 2024 Employer Information "
            "Release. U.S. Department of Homeland Security. Retrieved from "
            "https://www.uscis.gov/tools/reports-and-studies/"
            "h-1b-employer-data-hub"
        ),
        (
            "U.S. Citizenship and Immigration Services (USCIS). (2024). "
            "Characteristics of H-1B Specialty Occupation Workers: Fiscal "
            "Year 2024 Annual Report to Congress. U.S. Department of "
            "Homeland Security. Retrieved from "
            "https://www.uscis.gov/sites/default/files/document/reports/"
            "ola_signed_h1b_characteristics_congressional_report_FY24.pdf"
        ),
        (
            "U.S. Citizenship and Immigration Services (USCIS). (2024). "
            "Fiscal Year 2024: H-1B Petitions Annual Report to Congress. "
            "U.S. Department of Homeland Security. Retrieved from "
            "https://www.uscis.gov/sites/default/files/document/legal-docs/"
            "ola_signed_fy2024_h1b_petitions.pdf"
        ),
        (
            "U.S. Department of Labor, Office of Foreign Labor Certification "
            "(OFLC). (2024). LCA Disclosure Data — Labor Condition "
            "Application for H-1B, H-1B1, and E-3 Specialty Occupations "
            "(Performance Data). Employment and Training Administration. "
            "Retrieved from "
            "https://www.dol.gov/agencies/eta/foreign-labor/performance"
        ),
        (
            "U.S. Department of Labor, Office of Foreign Labor Certification "
            "(OFLC). (2024). Foreign Labor Application Gateway (FLAG) — "
            "Labor Condition Application Program Reference. Retrieved from "
            "https://flag.dol.gov/programs/LCA"
        ),
        (
            "U.S. Census Bureau. (2024). American Community Survey (ACS) "
            "5-Year Estimates, 2019–2023. U.S. Department of Commerce. "
            "Retrieved from https://www.census.gov/programs-surveys/acs/"
        ),
        (
            "U.S. Census Bureau. (2024). American Community Survey 5-Year "
            "Public Use Microdata Sample (PUMS), 2019–2023. Retrieved from "
            "https://www.census.gov/data/developers/data-sets/acs-5year.html"
        ),
        (
            "U.S. Bureau of Labor Statistics (BLS). (2024). Occupational "
            "Employment and Wage Statistics (OEWS). U.S. Department of "
            "Labor. Retrieved from https://www.bls.gov/oes/"
        ),
    ],
    "2. Methodology and Software": [
        (
            "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), "
            "5–32. https://doi.org/10.1023/A:1010933404324"
        ),
        (
            "Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree "
            "Boosting System. In Proceedings of the 22nd ACM SIGKDD "
            "International Conference on Knowledge Discovery and Data "
            "Mining (pp. 785–794). Association for Computing Machinery. "
            "https://doi.org/10.1145/2939672.2939785"
        ),
        (
            "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. "
            "(2002). SMOTE: Synthetic Minority Over-sampling Technique. "
            "Journal of Artificial Intelligence Research, 16, 321–357. "
            "https://doi.org/10.1613/jair.953"
        ),
        (
            "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., "
            "Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., "
            "Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., "
            "Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, E. "
            "(2011). Scikit-learn: Machine Learning in Python. Journal of "
            "Machine Learning Research, 12, 2825–2830."
        ),
        (
            "McKinney, W. (2010). Data Structures for Statistical Computing "
            "in Python. In Proceedings of the 9th Python in Science "
            "Conference (pp. 56–61). https://doi.org/10.25080/Majora-92bf1922-00a"
        ),
        (
            "Harris, C. R., Millman, K. J., van der Walt, S. J., Gommers, R., "
            "Virtanen, P., Cournapeau, D., et al. (2020). Array programming "
            "with NumPy. Nature, 585(7825), 357–362. "
            "https://doi.org/10.1038/s41586-020-2649-2"
        ),
        (
            "Zaharia, M., Xin, R. S., Wendell, P., Das, T., Armbrust, M., "
            "Dave, A., Meng, X., Rosen, J., Venkataraman, S., Franklin, M. J., "
            "Ghodsi, A., Gonzalez, J., Shenker, S., & Stoica, I. (2016). "
            "Apache Spark: A Unified Engine for Big Data Processing. "
            "Communications of the ACM, 59(11), 56–65. "
            "https://doi.org/10.1145/2934664"
        ),
        (
            "Vohra, D. (2016). Apache Parquet. In Practical Hadoop Ecosystem: "
            "A Definitive Guide to Hadoop-Related Frameworks and Tools "
            "(pp. 325–335). Apress. "
            "https://doi.org/10.1007/978-1-4842-2199-0_8"
        ),
        (
            "Apache Software Foundation. (2024). Apache Parquet: Columnar "
            "Storage Format. Retrieved from https://parquet.apache.org/"
        ),
        (
            "Apache Software Foundation. (2024). PySpark 3.5 Documentation. "
            "Retrieved from "
            "https://spark.apache.org/docs/3.5.0/api/python/index.html"
        ),
        (
            "Hastie, T., Tibshirani, R., & Friedman, J. (2009). The Elements "
            "of Statistical Learning: Data Mining, Inference, and Prediction "
            "(2nd ed.). Springer. https://doi.org/10.1007/978-0-387-84858-7"
        ),
        (
            "James, G., Witten, D., Hastie, T., & Tibshirani, R. (2021). An "
            "Introduction to Statistical Learning with Applications in R "
            "(2nd ed.). Springer. https://doi.org/10.1007/978-1-0716-1418-1"
        ),
    ],
    "3. H-1B Policy and Labor-Market Literature": [
        (
            "Kerr, W. R., & Lincoln, W. F. (2010). The Supply Side of "
            "Innovation: H-1B Visa Reforms and U.S. Ethnic Invention. "
            "Journal of Labor Economics, 28(3), 473–508. "
            "https://doi.org/10.1086/651934"
        ),
        (
            "Doran, K., Gelber, A., & Isen, A. (2022). The Effects of "
            "High-Skilled Immigration Policy on Firms: Evidence from "
            "Visa Lotteries. Journal of Political Economy, 130(10), "
            "2501–2533. https://doi.org/10.1086/720464"
        ),
        (
            "Dimmock, S. G., Huang, J., & Weisbenner, S. J. (2024). The "
            "Impact of Immigration on Firms and Workers: Insights from the "
            "H-1B Lottery. Federal Reserve Bank of Richmond Working Paper "
            "WP 24-04. Retrieved from "
            "https://www.richmondfed.org/-/media/RichmondFedOrg/"
            "publications/research/working_papers/2024/wp24-04.pdf"
        ),
        (
            "Peri, G., Shih, K., & Sparber, C. (2015). STEM Workers, H-1B "
            "Visas, and Productivity in U.S. Cities. Journal of Labor "
            "Economics, 33(S1), S225–S255. https://doi.org/10.1086/679061"
        ),
        (
            "Hunt, J., & Gauthier-Loiselle, M. (2010). How Much Does "
            "Immigration Boost Innovation? American Economic Journal: "
            "Macroeconomics, 2(2), 31–56. "
            "https://doi.org/10.1257/mac.2.2.31"
        ),
        (
            "Mayda, A. M., Ortega, F., Peri, G., Shih, K., & Sparber, C. "
            "(2018). The Effect of the H-1B Quota on the Employment and "
            "Selection of Foreign-Born Labor. European Economic Review, "
            "108, 105–128. https://doi.org/10.1016/j.euroecorev.2018.06.010"
        ),
        (
            "Bound, J., Khanna, G., & Morales, N. (2017). Understanding "
            "the Economic Impact of the H-1B Program on the U.S. NBER "
            "Working Paper No. 23153. National Bureau of Economic Research. "
            "https://doi.org/10.3386/w23153"
        ),
        (
            "Glennon, B. (2024). How do restrictions on high-skilled "
            "immigration affect offshoring? Evidence from the H-1B program. "
            "Management Science, 70(2), 907–930. "
            "https://doi.org/10.1287/mnsc.2023.4715"
        ),
    ],
    "4. Government & Think-Tank Reports": [
        (
            "U.S. Government Accountability Office (GAO). (2011). H-1B "
            "Visa Program: Reforms Are Needed to Minimize the Risks and "
            "Costs of Current Program (GAO-11-26). Retrieved from "
            "https://www.gao.gov/products/gao-11-26"
        ),
        (
            "U.S. Government Accountability Office (GAO). (2006). H-1B "
            "Visa Program: Labor Could Improve Its Oversight and Increase "
            "Information Sharing with Homeland Security (GAO-06-720). "
            "Retrieved from https://www.gao.gov/products/gao-06-720"
        ),
        (
            "Bier, D. J. (2025). H-1B Approvals Have Fallen Sharply Under "
            "New Restrictions. Cato Institute Policy Analysis. Cato "
            "Institute. Retrieved from https://www.cato.org/"
        ),
        (
            "National Foundation for American Policy (NFAP). (2018). "
            "H-1B Denials and Requests for Evidence Increase, Particularly "
            "for IT Services Companies. NFAP Policy Brief, July 2018. "
            "Retrieved from https://nfap.com/wp-content/uploads/2018/07/"
            "H-1B-Denial-and-RFE-Increase.NFAP-Policy-Brief.July-2018.pdf"
        ),
        (
            "National Foundation for American Policy (NFAP). (2025). "
            "H-1B Is the Most Restrictive Visa for High-Skilled "
            "Professionals. NFAP Policy Brief, March 2025. Retrieved from "
            "https://nfap.com/wp-content/uploads/2025/03/"
            "H-1B-Is-The-Most-Restrictive-Visa.NFAP-Policy-Brief.2025.pdf"
        ),
        (
            "Congressional Research Service (CRS). (2024). The H-1B Visa "
            "for Specialty Occupation Workers (IF12912). Library of "
            "Congress. Retrieved from "
            "https://www.congress.gov/crs-product/IF12912"
        ),
        (
            "Congressional Research Service (CRS). (2023). U.S. "
            "Employment-Based Immigration Policy (R47164). Library of "
            "Congress. Retrieved from "
            "https://www.congress.gov/crs-product/R47164"
        ),
        (
            "Migration Policy Institute (MPI). (2023). The U.S. "
            "High-Skilled Immigration System: Trends, Outcomes, and Policy "
            "Levers. Migration Policy Institute. Retrieved from "
            "https://www.migrationpolicy.org/"
        ),
        (
            "American Immigration Council. (2024). The H-1B Visa Program: "
            "A Primer on the Program and Its Impact on Jobs, Wages, and "
            "the Economy. Retrieved from "
            "https://www.americanimmigrationcouncil.org/"
        ),
    ],
    "5. Classification Systems and Technical Standards": [
        (
            "U.S. Census Bureau. (2022). North American Industry "
            "Classification System (NAICS) United States, 2022. Executive "
            "Office of the President, Office of Management and Budget. "
            "Retrieved from "
            "https://www.census.gov/naics/reference_files_tools/"
            "2022_NAICS_Manual.pdf"
        ),
        (
            "U.S. Bureau of Labor Statistics (BLS). (2018). Standard "
            "Occupational Classification (SOC) System, 2018. U.S. "
            "Department of Labor. Retrieved from "
            "https://www.bls.gov/soc/"
        ),
        (
            "U.S. Census Bureau. (2024). Census Regions and Divisions of "
            "the United States. U.S. Department of Commerce. Retrieved "
            "from https://www2.census.gov/geo/pdfs/maps-data/maps/"
            "reference/us_regdiv.pdf"
        ),
        (
            "U.S. Department of Homeland Security. (2024). Form I-129, "
            "Petition for a Nonimmigrant Worker — Instructions and "
            "Supporting Documentation. U.S. Citizenship and Immigration "
            "Services. Retrieved from https://www.uscis.gov/i-129"
        ),
        (
            "U.S. Department of Homeland Security. (2024). Form G-28, "
            "Notice of Entry of Appearance as Attorney or Accredited "
            "Representative. U.S. Citizenship and Immigration Services. "
            "Retrieved from https://www.uscis.gov/g-28"
        ),
    ],
}


def add_hyperlink_text(paragraph, text, font_size=11):
    """Add a normal text run to the paragraph (URLs in the citation text
    are kept as plain text — Word will auto-link them on open in most
    cases; this keeps the doc portable and free of hyperlink XML hacks).
    """
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"


def main():
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Default style: Times New Roman 11
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("References")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        "Building a Scalable Policy Monitoring Framework for "
        "H-1B Visa Approvals"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.name = "Times New Roman"

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_run = authors.add_run("Xiao Xu · Cassie Zhang")
    auth_run.font.size = Pt(11)
    auth_run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacer

    # Intro note
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "This document compiles the references cited and reasonably "
        "consultable for the H-1B explainability paper. Entries are "
        "grouped by category: primary data sources, methodology and "
        "software, H-1B policy and labor-market literature, government "
        "and think-tank reports, and classification systems. Citations "
        "follow APA 7 style; author / year / title / venue / DOI or URL "
        "are provided where available."
    )
    intro_run.font.size = Pt(11)
    intro_run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacer

    # Sections
    for section_title, refs in REFERENCES.items():
        h = doc.add_paragraph()
        h_run = h.add_run(section_title)
        h_run.bold = True
        h_run.font.size = Pt(13)
        h_run.font.name = "Times New Roman"

        for ref in refs:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(6)
            add_hyperlink_text(p, ref, font_size=11)

        doc.add_paragraph()  # spacer between sections

    # Closing note about how the references map to paper sections
    closing = doc.add_paragraph()
    c_run = closing.add_run("Notes on use")
    c_run.bold = True
    c_run.font.size = Pt(13)
    c_run.font.name = "Times New Roman"

    notes = [
        (
            "Section 1 (Primary Data Sources) covers the three datasets "
            "actually loaded by the pipeline (USCIS Employer Data Hub, "
            "OFLC LCA disclosure file, Census ACS) plus directly related "
            "USCIS congressional reports used to anchor descriptive "
            "statistics in Sections 1, 3, and 5 of the paper."
        ),
        (
            "Section 2 (Methodology and Software) supports Section 4 of "
            "the paper. Cite Breiman (2001) and Chen & Guestrin (2016) "
            "alongside the model names; Pedregosa et al. (2011) and "
            "McKinney (2010) cover the scikit-learn / pandas stack; "
            "Zaharia et al. (2016) and Vohra (2016) cover the PySpark "
            "and Parquet scalability story."
        ),
        (
            "Section 3 (H-1B Literature) supports the Introduction and "
            "Policy Implications sections. Kerr & Lincoln (2010), Doran "
            "et al. (2022), and Dimmock et al. (2024) are the strongest "
            "anchors for H-1B's labor-market role; Mayda et al. (2018) "
            "and Bound et al. (2017) cover the quota and macro effects."
        ),
        (
            "Section 4 (Government & Think-Tank) supports the policy "
            "framing: GAO and CRS reports establish that adjudication "
            "oversight gaps have been documented for two decades; NFAP "
            "and Cato briefs supply the recent denial / RFE rate "
            "context the paper references."
        ),
        (
            "Section 5 (Classification Systems) supports Section 3 (Data "
            "and Infrastructure) and the feature-engineering steps that "
            "rely on NAICS codes, SOC codes, and the I-129 / G-28 forms "
            "named in the Future Work section."
        ),
    ]

    for n in notes:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("•  " + n)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    # Footer with count summary
    doc.add_paragraph()
    total = sum(len(v) for v in REFERENCES.values())
    summary = doc.add_paragraph()
    s_run = summary.add_run(
        f"Total entries: {total} across {len(REFERENCES)} categories."
    )
    s_run.italic = True
    s_run.font.size = Pt(10)
    s_run.font.name = "Times New Roman"
    s_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")
    print(f"Total references: {total}")


if __name__ == "__main__":
    main()
