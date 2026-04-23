"""Build a Word document of the presentation script for slides 6-15."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "presentation_script_slides_6_to_15.docx"

# ---- content blocks: (heading, paragraphs) ----
SLIDES = [
    (
        "Slide 6 — End-to-End Reproducible Pipeline  (~45s)",
        [
            "So with the dataset introduced, I'll walk you through how the project is actually built. "
            "The pipeline has five stages, and the design rule is simple: one input, one output per stage, "
            "with Parquet at every boundary. Each script reads exactly one file and writes exactly one file — "
            "so any stage can be re-run on its own without rebuilding the whole thing.",

            "Following the chain on the left: we start from the raw USCIS CSV, then cleaning, feature engineering, "
            "the merge with external data, modeling, and analysis. By the end, sixty-one thousand employer rows "
            "become 3,968 merged cities.",

            "Three things to flag on the right. Storage is Parquet everywhere — columnar, compressed, schema-typed — "
            "so a full decade of releases would still fit on a laptop. Modularity means edit one stage, re-run only "
            "that stage. And orchestration is a thin shell: today it's a single Python script, but swapping to "
            "Airflow or Prefect later only changes the trigger, not the stages themselves.",
        ],
    ),
    (
        "Slide 7 — Honest Features from a Messy Public Release  (~50s)",
        [
            "A quick word on cleaning, because the public release is messier than it looks. The file ships as "
            "UTF-16 with a BOM. NAICS arrives as one mashed-up string we had to split into a code and a label. "
            "We added a zip3 column as a regional proxy. And critically — we re-derived the approval rate ourselves, "
            "because the raw column nudges values away from clean 0 and 1 by tiny amounts that distort group means.",

            "On the right are the engineered features. The structural set in green is everything we can build from "
            "USCIS alone: a STEM flag, a Census region, a size bucket, and log-volume controls. The enrichment set "
            "in yellow only exists at the city level after the merge: log wage from OFLC LCA, plus log population, "
            "density, rent, employment percentage, and the unemployment z-score. The whole feature set is "
            "deliberately bounded to what these public sources actually expose — zero leakage from the target.",
        ],
    ),
    (
        "Slide 8 — Same Merge, Pandas and PySpark  (~50s)",
        [
            "Now scalability. Today's data fits in memory, so the default path is pandas. But to show the pipeline "
            "is cluster-ready, we re-implemented the merge stage in PySpark 3.5. The code on the left is the actual "
            "excerpt — a SparkSession with adaptive execution turned on, a groupBy on state and city, and an inner "
            "join with the external file.",

            "The important part is on the right. We ran both versions on identical inputs and verified parity: same "
            "3,968 output rows, zero delta on approval rate, and an average-wage delta of less than 8 × 10⁻¹³ — "
            "pure float roundoff. So the Spark version is bit-identical to pandas.",

            "And the deploy story is one line. Locally we materialize through .toPandas().to_parquet. To go to "
            "Dataproc, EMR, or Databricks, you swap that for .write.parquet. Same DataFrame code, managed cluster, "
            "unchanged.",
        ],
    ),
    (
        "Slide 9 — Three Model Families, One Ceiling Question  (~50s)",
        [
            "For modeling, the goal is to bound explanatory power, not chase hyperparameters. The preprocessing "
            "pipeline on the left: numeric features pass through, categoricals go through a sparse one-hot with "
            "min_frequency=20 to keep the matrix manageable. Everything is wrapped in a sklearn Pipeline, an 80/20 "
            "split with a fixed seed so every model sees identical rows, and we score R² and MAE on the held-out "
            "test set.",

            "Three model families: OLS as the canonical baseline, a Random Forest, and XGBoost. And in the "
            "bottom-right, the key design choice — two specs run on identical rows. Structural uses only "
            "USCIS-derivable features. Enriched adds the wage and local-economy variables. The delta between them "
            "is the test of whether public enrichment actually moves the ceiling.",
        ],
    ),
    (
        "Slide 10 — Layer 1: Demand (Cassie's results)  (~30s)",
        [
            "Layer one was contributed by my partner Cassie. Her question: can public economic geography predict "
            "where H-1B petitions cluster? She bucketed roughly 9,800 cities into low, medium, and high petition "
            "volume, and ran XGBoost with SMOTE plus a state feature. The model hits an AUC of 0.79 on the "
            "high-demand class, with a weighted F1 around 0.60.",

            "So wages and density do carry real signal — California, New Jersey, and Texas concentrate the "
            "high-demand cities. But — and this is the bridge to the rest of my part — predicting where filings "
            "happen is not the same as predicting whether they get approved.",
        ],
    ),
    (
        "Slide 11 — Layer 2: USCIS Alone Explains Essentially Nothing  (~50s)",
        [
            "Layer two is the baseline approval question on USCIS data alone. Three model families, same near-zero "
            "ceiling. OLS, Random Forest, and XGBoost all land at R² between 0.006 and 0.008, with MAE around six "
            "percentage points.",

            "Why? Two reasons. First, 82% of rows already sit at 99.9% approval or higher — the target is "
            "near-degenerate. Second, look at the between-group variance table on the right: state explains 0.33% "
            "of variance, region 0.06%, NAICS industry 0.30%, size bucket 0.03%. Every observable grouping accounts "
            "for less than half a percent. There is essentially nothing here to model.",
        ],
    ),
    (
        "Slide 12 — Layer 3: Wages and Local Economy Still Can't Beat the Mean  (~50s)",
        [
            "Layer three is the real test. Aggregate USCIS to the city level, inner-join Cassie's enriched file, "
            "run identical rows on both specs.",

            "Read this table carefully. The structural specs all sit below zero — that's tree models overfitting "
            "noise. When we add the enrichment columns, tree R² rises sharply: Random Forest gains 17 points, "
            "XGBoost gains 14. But none of them cross zero. OLS only improves by 0.007. The MAE settles around 0.04 "
            "— about four percentage points of noise around a mean of 96.6%. The headline: a constant-mean predictor "
            "beats every model we tested.",
        ],
    ),
    (
        "Slide 13 — Interpretation: The Data Is Bounded, Not the Method  (~45s)",
        [
            "Three layers, one converging answer. Layer 1: where H-1B workers are needed is publicly visible. "
            "Layer 2: at the employer level, approval is not — observable groupings explain under half a percent. "
            "Layer 3: even after we add wage and local-economy enrichment across nearly 4,000 cities, R² stays at "
            "or below zero.",

            "So the conclusion is sharper than \"our models are bad.\" Approval variance lives inside individual "
            "case files — wage relative to prevailing wage by SOC, RFE history, the attorney, documentation "
            "quality. None of those signals exist in any public release.",
        ],
    ),
    (
        "Slide 14 — Limitations  (~30s)",
        [
            "To be explicit about scope: we make no predictive claim — the near-zero R² is the finding. We make no "
            "causal claim — state and industry correlate with approval but don't cause it. We observe only the "
            "post-lottery 25% who reached adjudication. And FY2024 alone can't separate structural patterns from "
            "administration-specific policy regimes.",
        ],
    ),
    (
        "Slide 15 — Future Work  (~30s)",
        [
            "The bottleneck was never compute — it was the input. So the commit worth making is data, not hardware. "
            "Three priorities. First, case-level data — joining USCIS I-129 releases with wage-relative-to-prevailing-by-SOC, "
            "RFE history, and attorney metadata. That's where the unexplained variance lives. Second, multi-year "
            "coverage from FY2019 through FY2025 to separate policy-regime effects from structural ones. Third, "
            "productionize the Spark pipeline — that one-line swap to a managed cluster.",

            "Bounding what public data can't explain is the first step to demanding better data. Thank you.",
        ],
    ),
]

DELIVERY_TIPS = [
    "Pace target: ~150 words/min lands you near 7 minutes; slow to ~135 wpm for closer to 8.",
    "Slow down on the numbers in slides 11 and 12 — the audience needs a beat to read the tables.",
    "Slides 14–15 are short on purpose — they're closing beats, not new content. Don't rush slide 13's "
    "\"approval variance lives inside case files\" line; that's your thesis.",
]


def main():
    doc = Document()

    # Page margins — comfortable for printing / reading on screen.
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    # Default body style — Calibri 11
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Presentation Script — Slides 6 to 15")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("H-1B Approval: How Much Can Public Data Explain?  ·  ~7 minutes  ·  ~1,150 words")
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph()  # spacer

    # Per-slide blocks
    for heading, paragraphs in SLIDES:
        h = doc.add_paragraph()
        h_run = h.add_run(heading)
        h_run.bold = True
        h_run.font.size = Pt(13)
        h_run.font.color.rgb = RGBColor(0x1C, 0x72, 0x93)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)

        for para in paragraphs:
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.35

    # Delivery tips section
    doc.add_paragraph()
    tips_h = doc.add_paragraph()
    tips_run = tips_h.add_run("Delivery Tips")
    tips_run.bold = True
    tips_run.font.size = Pt(13)
    tips_run.font.color.rgb = RGBColor(0x1C, 0x72, 0x93)
    tips_h.paragraph_format.space_before = Pt(14)
    tips_h.paragraph_format.space_after = Pt(4)

    for tip in DELIVERY_TIPS:
        p = doc.add_paragraph(tip, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.35

    doc.save(str(OUT))
    print(f"saved {OUT}")


if __name__ == "__main__":
    main()
